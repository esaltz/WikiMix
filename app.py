# Import necessary libraries
from bs4 import BeautifulSoup
from docx import Document
from email.message import EmailMessage
from fpdf import FPDF
from gspread_dataframe import get_as_dataframe, set_with_dataframe
from IPython.display import display, HTML, IFrame
from matplotlib.backends.backend_pdf import PdfPages
from nltk.tokenize import sent_tokenize
from PIL import Image
from st_aggrid import GridOptionsBuilder, AgGrid, GridUpdateMode, DataReturnMode
from streamlit.components.v1 import html
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

import certifi
import datetime
import gspread
import helper
import hydralit_components as hc
import json
import joblib
import matplotlib.pyplot as plt
import nltk
import numpy as np
import openpyxl
import os
import pandas as pd
import pprint
import pybase64
import random
import re
import requests
import smtplib
import spacy
import sqlite3
import ssl 
ssl._create_default_https_context = ssl._create_unverified_context

import streamlit as st
import streamlit.components.v1 as components
import string
import subprocess
import sys
import tensorflow as tf
import time
import urllib.request
import warnings
import wikipedia 
import wikipediaapi

#make it look nice from the start
st.set_page_config(layout='wide',page_title='WikiMix')

#Hide menu and footer
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            [data-testid="stForm"] {border: 0px; align-items: center}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True) 

# Warnings ignore 
warnings.filterwarnings(action='ignore')
st.set_option('deprecation.showfileUploaderEncoding', False)
st.set_option('deprecation.showPyplotGlobalUse', False)

sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
sheet_name = "info_data"
sheet2_name = "data"
cred_file = "creds.json"
gc = gspread.service_account(cred_file)
database = gc.open("WikiMix")
wks_info_data = database.worksheet(sheet_name)
wks_data = database.worksheet(sheet2_name)

def main():
    # specify the primary menu definition
    menu_data = [
        {'icon': "fa fa-newspaper", 'label':"Curate a WikiMix"},
        {'icon': "fa fa-envelope", 'label':"Deliver a WikiMix"},
        {'icon': "fa fa-search", 'label':"Search for WikiMix"},
        {'icon': "fa fa-question", 'label':"Random WikiMix"},
        {'icon': "fa fa-database", 'label':"Explore WikiMix Database"},
        {'icon': "fa fa-database", 'label':"Test"}
        ]

    over_theme = {'txc_inactive': '#fffae1', 'menu_background':'#3366CC'}
    menu_id = hc.nav_bar(
        menu_definition=menu_data,
        override_theme=over_theme,
        home_name='WikiMix',
        hide_streamlit_markers=False, #will show the st hamburger as well as the navbar now!
        sticky_nav=False, #at the top or not
        sticky_mode='pinned', #jumpy or not-jumpy, but sticky or pinned
    )
    
    #What is WikiMix?
    if menu_id == 'WikiMix':
        st.subheader("What is WikiMix?")
        st.write("WikiMix is an exciting game that takes you on a journey through the vast knowledge of Wikipedia! In this game, you'll create a unique list of Wikipedia articles, known as a WikiMix. There are mulitple ways to create a WikiMix.")
        st.markdown(
        """
        - **Curate a WikiMix** - Design a mix of your favorite Wikis using only your know how
        - **Deliver a WikiMix** - Give a mix of your favorite Wikis to someone special
        - **Search WikiMix** - Find Wikis you know nothing about using subjects, themes, or tags
        - **Random WikiMix** - Press a button and have a generated mix produced for you
        - **Explore WikiMix Database** - Search through previous WikiMixes
        """
        )

        st.markdown("---")
        
        st.subheader("How-to Guide")
        st.write("Creating a WikiMix is easy and simple.")
        st.write("Enter in Information. Each WikiMix requires a Title. Also, other information you may want to include are the following: a Note about the mix; who the mix is for; who the mix is made by; and number of articles included in the mix")
        st.write("Next, enter article titles and links to add to the mix by pressing the 'Add Article' button. When you have finished adding articles move to 'The Mix' section")
        st.write("Most importantly, Have fun creating WikiMixes.")

    #Curate a Mix
    elif menu_id == "Curate a WikiMix":
        st.header("Curate a WikiMix")
        with st.container():
            sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
            sheet_name = "info_data"
            sheet2_name = "data"
            cred_file = "creds.json"
            gc = gspread.service_account(cred_file)
            database = gc.open("WikiMix")
            wks_info_data = database.worksheet(sheet_name)
            wks_data = database.worksheet(sheet2_name)
            wikiTitle = st.text_input("Title of WikiMix")
            if st.button('Submit WikiTitle'):
                helper.update_cell_wikiTitle(wikiTitle)
            def generate_new_wiki(wikiTitle):
                if wikiTitle != wks_info_data.col_values(1):
                    sheet3_name = (wikiTitle + " WikiMix List")
                    database.add_worksheet(title=sheet3_name, rows=100, cols=20)
            col1,col2 = st.columns([1.5,1.5])
            def update_cell_TitleMix_wikiFor(wikiFor):
                next_empty_row = helper.find_next_empty_cell_wikiFor_column(5)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 5, wikiFor)

            def update_cell_TitleMix_wikiBy(wikiBy):
                next_empty_row = helper.find_next_empty_cell_wikiBy_column(6)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 6, wikiBy)

            def update_cell_TitleMix_wikiArticleTitle(article):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleTitle_column(1)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 1, article)

            def update_cell_TitleMix_wikiArticleLink(url):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleLink_column(2)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 2, url)

            def update_cell_TitleMix_wikiArticleSummary(summary):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleSummary_column(3)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 3, summary)

            def wikipedia_search(wikiArticleTitle):
                try:
                    page = wikipedia.page(wikiArticleTitle)
                    url = page.url
                    summary = page.summary
                    summary_sentences = sent_tokenize(summary)
                    summarized_summary = ' '.join(summary_sentences[:3])
                    return url, summarized_summary
                except wikipedia.exceptions.DisambiguationError as e:
                    options = e.options  # List of disambiguation options
                    first_option = options[0]  # Choose the first option automatically
                    try:
                        page = wikipedia.page(first_option)
                        url = page.url
                        summary = page.summary
                        summary_sentences = sent_tokenize(summary)
                        summarized_summary = ' '.join(summary_sentences[:3])
                        return url, summarized_summary
                    except wikipedia.exceptions.PageError:
                        st.warning(f"No Wikipedia page found for '{first_option}'")
                        return None, None
                except wikipedia.exceptions.PageError:
                    # If the page does not exist, return None for URL and summary
                    return None, None

            with col1:
                with st.form(key='info_form', clear_on_submit=True):
                    st.subheader("Info Gathering")
                    wikiNote = st.text_area("Note", height=208)
                    submitted = st.form_submit_button("Info Gathering")
                    sheet3_name = (wikiTitle + " WikiMix List")
                    if submitted:
                        helper.update_cell_wikiNote(wikiNote)
                        generate_new_wiki(wikiTitle)
                        wks_wikiTitle_data = database.worksheet(sheet3_name)
                        wks_wikiTitle_data.format('A1:F1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(1,4, 'WikiMix Note')
                        wks_wikiTitle_data.update_cell(2,4, wikiNote)

            with col2:
                with st.form(key='article_form', clear_on_submit=True):
                    st.subheader("Add Articles")
                    wikiArticleTitle = st.text_input("Title of Article")
                    wikiArticleLink = st.text_input("Web address")
                    submitted2a = st.form_submit_button("Add Article")
                    if submitted2a:
                        st.success("Article Added to WikiMix. Go back and add another!")
                        sheet3_name = (wikiTitle + " WikiMix List")
                        wks_wikiTitle_data = database.worksheet(sheet3_name)
                        wks_wikiTitle_data.format('A1:D1', {'textFormat': {'bold': True}})

                        wks_wikiTitle_data.update_cell(1,1, 'Article Title')
                        wks_wikiTitle_data.update_cell(1,2, 'Article Link')
                        wks_wikiTitle_data.update_cell(1,3, 'Article Summary')

                        helper.update_cell_wikiArticleTitle(wikiArticleTitle)
                        helper.update_cell_wikiArticleLink(wikiArticleLink)

                        #TitleMix data
                        def find_next_empty_cell_TitleMix_wikiArticleTitle_column(column):
                            # Get all non-empty values in the specified column
                            values = wks_wikiTitle_data.col_values(1)

                            # Find the index of the first empty cell
                            for idx, value in enumerate(values):
                                if not value:
                                    return idx + 1  # Adding 1 to convert to 1-based index

                            # If no empty cell found, return the next index after the last cell
                            return len(values) + 1

                        def find_next_empty_cell_TitleMix_wikiArticleLink_column(column):
                            # Get all non-empty values in the specified column
                            values = wks_wikiTitle_data.col_values(2)

                            # Find the index of the first empty cell
                            for idx, value in enumerate(values):
                                if not value:
                                    return idx + 1  # Adding 1 to convert to 1-based index

                            # If no empty cell found, return the next index after the last cell
                            return len(values) + 1

                        def find_next_empty_cell_TitleMix_wikiArticleSummary_column(column):
                            # Get all non-empty values in the specified column
                            values = wks_wikiTitle_data.col_values(3)

                            # Find the index of the first empty cell
                            for idx, value in enumerate(values):
                                if not value:
                                    return idx + 1  # Adding 1 to convert to 1-based index

                            # If no empty cell found, return the next index after the last cell
                            return len(values) + 1
             
                        article_list = wks_data.col_values(1)[1:]
                        nltk.download('punkt')
                        latest_article = None
                        latest_url = None
                        latest_summary = None

                        for article in article_list:
                            url, summary = helper.wikipedia_search(article)  # Pass the article as an argument
                            if url and summary:
                                latest_article = article
                                latest_url = url
                                latest_summary = summary

                        # Display the latest updated information
                        if latest_article:
                            st.markdown(f"**Title:** {latest_article}")
                            st.markdown(f"**URL:** [Link]({latest_url})")
                            st.markdown(f"**Summary:** {latest_summary}")
                        #else:
                            #st.warning("No Wikipedia page found for the latest article.")

                            def find_next_empty_cell_wikiArticleSummary_column(column):
                                # Get all non-empty values in the specified column
                                values = wks_data.col_values(3)

                                # Find the index of the first empty cell
                                for idx, value in enumerate(values):
                                    if not value:
                                        return idx + 1  # Adding 1 to convert to 1-based index

                                # If no empty cell found, return the next index after the last cell
                                return len(values) + 1

                            def update_cell_wikiArticleSummary(latest_summary):
                                next_empty_row = find_next_empty_cell_wikiArticleSummary_column(3)  # Column A is 1-based index
                                wks_data.update_cell(next_empty_row, 3, latest_summary)

                            update_cell_wikiArticleSummary(latest_summary)

                            update_cell_TitleMix_wikiArticleTitle(wikiArticleTitle)
                            update_cell_TitleMix_wikiArticleLink(wikiArticleLink)
                            update_cell_TitleMix_wikiArticleSummary(latest_summary)

            if wikiTitle and wikiFor and wikiBy and wikiNote:
                st.markdown("---")
                st.subheader("The Mix")
                try:
                    sheet3_name = wikiTitle + " WikiMix List"
                    wks_wikiTitle_data = database.worksheet(sheet3_name)
                    article_list = wks_wikiTitle_data.col_values(1)[1:]
                    wikiFor = wks_wikiTitle_data.acell('E2').value
                    wikiBy = wks_wikiTitle_data.acell('F2').value
                    wikiNote = wks_wikiTitle_data.acell('D2').value 

                    def create_pdf_document(article_list):
                        pdf = FPDF()
                        pdf.set_auto_page_break(auto=True, margin=15)
                        pdf.add_page()
                        
                        pdf.set_font("Arial", size=28)
                        pdf.cell(200, 10, txt=wikiTitle + " WikiMix", ln=True, align="C")

                        pdf.set_font("Arial", size=14, style="B")
                        pdf.cell(200, 10, txt = "Note: ", ln=True, align='L')
                        pdf.set_font("Arial", size=12)
                        pdf.multi_cell(0,5, txt = wikiNote, align='L')
                        pdf.ln()

                        for article in article_list:
                            url, summary = helper.wikipedia_search(article)
                            if url and summary:
                                pdf.set_font("Arial", size=14, style="B")
                                pdf.cell(200, 10, txt=article.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                                pdf.set_text_color(0, 0, 255)
                                pdf.set_font("Arial", size=10)
                                pdf.cell(0, 10, txt=url, ln=True, link=url)
                                pdf.set_font("Arial", size=12, style="B")
                                pdf.set_text_color(0, 0, 0)
                                pdf.cell(200, 10, txt="Summary:", ln=True)
                                pdf.set_font("Arial", size=12)
                                pdf.multi_cell(0, 10, txt=summary.encode('latin-1', 'replace').decode('latin-1'),  align='L')
                                pdf.ln()
                        
                        return pdf

                    def create_word_document(article_list):
                        doc = Document()
                        doc.add_heading(wikiTitle + " WikiMix", level=1)
                        doc.add_heading("Made For", level=2)
                        doc.add_paragraph(wikiFor)
                        doc.add_heading("Made By", level=2)
                        doc.add_paragraph(wikiBy)
                        doc.add_heading("Note:", level=2)
                        doc.add_paragraph(wikiNote)

                        bold_style = doc.styles.add_style('BoldStyle', WD_STYLE_TYPE.CHARACTER)
                        bold_font = bold_style.font
                        bold_font.bold = True
                        
                        for article in article_list:
                            url, summary = helper.wikipedia_search(article)
                            if url and summary:
                                doc.add_heading(article, level=2)
                                doc.add_heading(f"URL:", level=2).bold = True
                                doc.add_paragraph(f"{url}", style=None)
                                doc.add_heading(f"Summary:", level=2).bold = True
                                doc.add_paragraph(f"{summary}", style=None)
                        
                        return doc 

                    if st.button("See "+ wikiTitle + " WikiMix"):
                        st.write("WikiMix: " + wikiTitle)
                        st.write("WikiMix Note: " + wks_wikiTitle_data.acell('D2').value)

                        for article in article_list:
                            if article:  # Check if the cell is not empty
                                url, summary = helper.wikipedia_search(article)
                                if url and summary:
                                    st.markdown(f"**Title:** {article}")
                                    st.markdown(f"**URL:** [Link]({url})")
                                    st.markdown("**Summary:**")
                                    st.write(summary)
                                else:
                                    st.warning(f"No Wikipedia page found for '{article}'")
                            else:
                                break  # Break the loop when an empty cell is encountered

                        doc = create_word_document(article_list)
                        doc_filename = wikiTitle + " WikiMix.docx"
                        doc.save(doc_filename)
                        st.markdown(helper.get_download_link(doc_filename, "Download Doc"), unsafe_allow_html=True)
                        
                        pdf = create_pdf_document(article_list)
                        pdf_path = wikiTitle + " WikiMix.pdf"
                        create_pdf_document(article_list).output(pdf_path)
                        st.markdown(helper.get_download_link(pdf_path, "Download PDF"), unsafe_allow_html=True)

                    st.markdown("---")
                    # Email File
                    st.subheader("Send WikiMix")
                    email_sender = 'WikiMixList@gmail.com'
                    email_password = 'ufpkegvfsqciasab'
                    email_receiver = st.text_input("Email To")
                    subject = "Here's your WikiMix"
                    body = "Welcome to WikiMix, the app that let's you expertly curate a collection of Wikipedia articles that traverse diverse topics to offer an immersive journey through the realm of information." 
                    
                    pdf = create_pdf_document(article_list)
                    pdf_path = wikiTitle + " WikiMix.pdf"
                    create_pdf_document(article_list).output(pdf_path)
                    
                    em = EmailMessage()
                    em['From'] = email_sender
                    em['To'] = email_receiver
                    em['Subject'] = subject
                    em.set_content(body)

                    context = ssl.create_default_context()
                    button = st.button("Send WikiMix")
                    filename = pdf_path

                    if button:
                        with open(filename, 'rb') as f:
                            file_data = f.read()
                            em.add_attachment(file_data, maintype='application', subtype='pdf', filename=filename)
                        with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
                            smtp.login(email_sender, email_password)
                            smtp.sendmail(email_sender,email_receiver, em.as_string())
                        st.success("WikiMix has been sent!")
                except gspread.exceptions.WorksheetNotFound:
                    article_list = []

    #Deliver a Mix
    elif menu_id == "Deliver a WikiMix":
        st.header("Deliver a WikiMix")
        with st.container():
            sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
            sheet_name = "info_data"
            sheet2_name = "data"
            cred_file = "creds.json"
            gc = gspread.service_account(cred_file)
            database = gc.open("WikiMix")
            wks_info_data = database.worksheet(sheet_name)
            wks_data = database.worksheet(sheet2_name)
            wikiTitle = st.text_input("Title of WikiMix")
            if st.button('Submit WikiTitle'):
                helper.update_cell_wikiTitle(wikiTitle)
            
            def generate_new_wiki(wikiTitle):
                if wikiTitle != wks_info_data.col_values(1):
                    sheet3_name = (wikiTitle + " WikiMix List")
                    database.add_worksheet(title=sheet3_name, rows=100, cols=20)
            
            def update_cell_TitleMix_wikiFor(wikiFor):
                next_empty_row = helper.find_next_empty_cell_wikiFor_column(5)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 5, wikiFor)

            def update_cell_TitleMix_wikiBy(wikiBy):
                next_empty_row = helper.find_next_empty_cell_wikiBy_column(6)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 6, wikiBy)

            def update_cell_TitleMix_wikiArticleTitle(article):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleTitle_column(1)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 1, article)

            def update_cell_TitleMix_wikiArticleLink(url):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleLink_column(2)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 2, url)

            def update_cell_TitleMix_wikiArticleSummary(summary):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleSummary_column(3)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 3, summary)

            def wikipedia_search(wikiArticleTitle):
                try:
                    page = wikipedia.page(wikiArticleTitle)
                    url = page.url
                    summary = page.summary
                    summary_sentences = sent_tokenize(summary)
                    summarized_summary = ' '.join(summary_sentences[:3])
                    return url, summarized_summary
                except wikipedia.exceptions.DisambiguationError as e:
                    options = e.options  # List of disambiguation options
                    first_option = options[0]  # Choose the first option automatically
                    try:
                        page = wikipedia.page(first_option)
                        url = page.url
                        summary = page.summary
                        summary_sentences = sent_tokenize(summary)
                        summarized_summary = ' '.join(summary_sentences[:3])
                        return url, summarized_summary
                    except wikipedia.exceptions.PageError:
                        st.warning(f"No Wikipedia page found for '{first_option}'")
                        return None, None
                except wikipedia.exceptions.PageError:
                    # If the page does not exist, return None for URL and summary
                    return None, None

            col1,col2 = st.columns([2.5,2.5])
            with col1:
                with st.form(key='info_form', clear_on_submit=True):
                    st.subheader("Info Gathering")
                    wikiFor = st.text_input("Who is this WikiMix for?")
                    wikiBy = st.text_input("Who was this WikiMix made by?")
                    wikiNote = st.text_area("Note", height=122)
                    submitted = st.form_submit_button("Info Gathering")
                    sheet3_name = (wikiTitle + " WikiMix List")
                    if submitted:
                        st.subheader("The " + wikiTitle + " WikiMix")
                        st.write("Made for " + wikiFor)
                        st.write("Made by " + wikiBy)
                        st.write("Note: " + wikiNote)
                        helper.update_cell_wikiNote(wikiNote)
                        helper.update_cell_wikiFor(wikiFor)
                        helper.update_cell_wikiBy(wikiBy)
                        generate_new_wiki(wikiTitle)
                        
                        wks_wikiTitle_data = database.worksheet(sheet3_name)
                        wks_wikiTitle_data.format('A1:F1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(1,4, 'WikiMix Note')
                        wks_wikiTitle_data.update_cell(2,4, wikiNote)

                        wks_wikiTitle_data.update_cell(1,5, 'WikiMix For')
                        wks_wikiTitle_data.update_cell(2,5, wikiFor)

                        wks_wikiTitle_data.update_cell(1,6, 'WikiMix By')
                        wks_wikiTitle_data.update_cell(2,6, wikiBy)

            with col2:
                with st.form(key='article_form', clear_on_submit=True):
                    st.subheader("Add Articles")
                    wikiArticleTitle = st.text_input("Title of Article")
                    wikiArticleLink = st.text_area("Web address", height=210)

                    #TitleMix data
                    def find_next_empty_cell_TitleMix_wikiArticleTitle_column(column):
                        # Get all non-empty values in the specified column
                        values = wks_wikiTitle_data.col_values(1)

                        # Find the index of the first empty cell
                        for idx, value in enumerate(values):
                            if not value:
                                return idx + 1  # Adding 1 to convert to 1-based index

                        # If no empty cell found, return the next index after the last cell
                        return len(values) + 1

                    def find_next_empty_cell_TitleMix_wikiArticleLink_column(column):
                        # Get all non-empty values in the specified column
                        values = wks_wikiTitle_data.col_values(2)

                        # Find the index of the first empty cell
                        for idx, value in enumerate(values):
                            if not value:
                                return idx + 1  # Adding 1 to convert to 1-based index

                        # If no empty cell found, return the next index after the last cell
                        return len(values) + 1

                    def find_next_empty_cell_TitleMix_wikiArticleSummary_column(column):
                        # Get all non-empty values in the specified column
                        values = wks_wikiTitle_data.col_values(3)

                        # Find the index of the first empty cell
                        for idx, value in enumerate(values):
                            if not value:
                                return idx + 1  # Adding 1 to convert to 1-based index

                        # If no empty cell found, return the next index after the last cell
                        return len(values) + 1
  
                    def find_next_empty_cell_wikiArticleSummary_column(column):
                            # Get all non-empty values in the specified column
                            values = wks_data.col_values(3)

                            # Find the index of the first empty cell
                            for idx, value in enumerate(values):
                                if not value:
                                    return idx + 1  # Adding 1 to convert to 1-based index

                            # If no empty cell found, return the next index after the last cell
                            return len(values) + 1

                    def update_cell_wikiArticleSummary(latest_summary):
                        next_empty_row = find_next_empty_cell_wikiArticleSummary_column(3)  # Column A is 1-based index
                        wks_data.update_cell(next_empty_row, 3, latest_summary)

                    submitted2a = st.form_submit_button("Add Article")
                    if submitted2a:
                        st.success("Article Added to WikiMix. Go back and add another!")
                        sheet3_name = (wikiTitle + " WikiMix List")
                        wks_wikiTitle_data = database.worksheet(sheet3_name)
                        wks_wikiTitle_data.format('A1:F1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(1,1, 'Article Title')
                        wks_wikiTitle_data.update_cell(1,2, 'Article Link')
                        wks_wikiTitle_data.update_cell(1,3, 'Article Summary')
                        wks_wikiTitle_data.update_cell(1,5, 'WikiMix For')
                        wks_wikiTitle_data.update_cell(1,6, 'WikiMix By')
                        helper.update_cell_wikiArticleTitle(wikiArticleTitle)
                        helper.update_cell_wikiArticleLink(wikiArticleLink)
                        
                        article_list = wks_data.col_values(1)[1:]
                        #nltk.download('punkt')
                        latest_article = None
                        latest_url = None
                        latest_summary = None

                        for article in article_list:
                            url, summary = helper.wikipedia_search(article)  # Pass the article as an argument
                            if url and summary:
                                latest_article = article
                                latest_url = url
                                latest_summary = summary

                        # Display the latest updated information
                        if latest_article:
                            st.markdown(f"**Title:** {latest_article}")
                            st.markdown(f"**URL:** [Link]({latest_url})")
                            st.markdown(f"**Summary:** {latest_summary}")
                            
                            update_cell_wikiArticleSummary(latest_summary)
                            update_cell_TitleMix_wikiArticleTitle(wikiArticleTitle)
                            update_cell_TitleMix_wikiArticleLink(wikiArticleLink)
                            update_cell_TitleMix_wikiArticleSummary(latest_summary)

            if wikiTitle and wikiFor and wikiBy and wikiNote:
                st.markdown("---")
                st.subheader("The Mix")
                sheet3_name = wikiTitle + " WikiMix List"
                wks_wikiTitle_data = database.worksheet(sheet3_name)
                article_list = wks_wikiTitle_data.col_values(1)[1:]
                wikiFor = wks_wikiTitle_data.acell('E2').value
                wikiBy = wks_wikiTitle_data.acell('F2').value
                wikiNote = wks_wikiTitle_data.acell('D2').value 

                def create_pdf_document(article_list):
                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    
                    pdf.set_font("Arial", size=28)
                    pdf.cell(200, 10, txt=wikiTitle + " WikiMix", ln=True, align="C")
                    
                    pdf.set_font("Arial", size=14, style="B")
                    pdf.cell(200, 10, txt = "Made For ", ln=True, align='L')
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0,5, txt = wikiFor)

                    pdf.set_font("Arial", size=14, style="B")
                    pdf.cell(200, 10, txt = "Made By ", ln=True, align='L')
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0,5, txt = wikiBy)

                    pdf.set_font("Arial", size=14, style="B")
                    pdf.cell(200, 10, txt = "Note: ", ln=True, align='L')
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0,5, txt = wikiNote, align='L')
                    pdf.ln()

                    for article in article_list:
                        url, summary = helper.wikipedia_search(article)
                        if url and summary:
                            pdf.set_font("Arial", size=14, style="B")
                            pdf.cell(200, 10, txt=article.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                            pdf.set_text_color(0, 0, 255)
                            pdf.set_font("Arial", size=10)
                            pdf.cell(0, 10, txt=url, ln=True, link=url)
                            pdf.set_font("Arial", size=12, style="B")
                            pdf.set_text_color(0, 0, 0)
                            pdf.cell(200, 10, txt="Summary:", ln=True)
                            pdf.set_font("Arial", size=12)
                            pdf.multi_cell(0, 10, txt=summary.encode('latin-1', 'replace').decode('latin-1'),  align='L')
                            pdf.ln()
                    
                    return pdf

                def create_word_document(article_list):
                    doc = Document()
                    doc.add_heading(wikiTitle + " WikiMix", level=1)
                    doc.add_heading("Made For", level=2)
                    doc.add_paragraph(wikiFor)
                    doc.add_heading("Made By", level=2)
                    doc.add_paragraph(wikiBy)
                    doc.add_heading("Note:", level=2)
                    doc.add_paragraph(wikiNote)

                    bold_style = doc.styles.add_style('BoldStyle', WD_STYLE_TYPE.CHARACTER)
                    bold_font = bold_style.font
                    bold_font.bold = True
                    
                    for article in article_list:
                        url, summary = helper.wikipedia_search(article)
                        if url and summary:
                            doc.add_heading(article, level=2)
                            doc.add_heading(f"URL:", level=2).bold = True
                            doc.add_paragraph(f"{url}", style=None)
                            doc.add_heading(f"Summary:", level=2).bold = True
                            doc.add_paragraph(f"{summary}", style=None)
                    
                    return doc 

                if st.button("See "+ wikiTitle + " WikiMix"):
                    doc = create_word_document(article_list)
                    doc_filename = wikiTitle + " WikiMix.docx"
                    doc.save(doc_filename)
                    st.markdown(helper.get_download_link(doc_filename, "Download Doc"), unsafe_allow_html=True)
                    
                    pdf = create_pdf_document(article_list)
                    pdf_path = wikiTitle + " WikiMix.pdf"
                    create_pdf_document(article_list).output(pdf_path)
                    st.markdown(helper.get_download_link(pdf_path, "Download PDF"), unsafe_allow_html=True)

                # Email File
                st.subheader("Send WikiMix")
                email_sender = 'WikiMixList@gmail.com'
                email_password = 'ufpkegvfsqciasab'
                email_receiver = st.text_input("Email To")
                subject = "Here's your WikiMix"
                body = "Welcome to WikiMix, the app that let's you expertly curate a collection of Wikipedia articles that traverse diverse topics to offer an immersive journey through the realm of information." 
                
                pdf = create_pdf_document(article_list)
                pdf_path = wikiTitle + " WikiMix.pdf"
                create_pdf_document(article_list).output(pdf_path)
                
                em = EmailMessage()
                em['From'] = email_sender
                em['To'] = email_receiver
                em['Subject'] = subject
                em.set_content(body)

                context = ssl.create_default_context()
                button = st.button("Send WikiMix")
                filename = pdf_path

                if button:
                    with open(filename, 'rb') as f:
                        file_data = f.read()
                        em.add_attachment(file_data, maintype='application', subtype='pdf', filename=filename)
                    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
                        smtp.login(email_sender, email_password)
                        smtp.sendmail(email_sender,email_receiver, em.as_string())
                    st.success("WikiMix has been sent!")

    #Search for WikiMix
    elif menu_id == "Search for WikiMix":
        st.header("Search for WikiMix")
        with st.container():
            sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
            sheet_name = "info_data"
            sheet2_name = "data"
            cred_file = "creds.json"
            gc = gspread.service_account(cred_file)
            database = gc.open("WikiMix")
            wks_info_data = database.worksheet(sheet_name)
            wks_data = database.worksheet(sheet2_name)
            wikiTitle = st.text_input("Title of WikiMix")
            if st.button('Submit WikiTitle'):
                helper.update_cell_wikiTitle(wikiTitle)
            def generate_new_wiki(wikiTitle):
                if wikiTitle != wks_info_data.col_values(1):
                    sheet3_name = (wikiTitle + " WikiMix List")
                    database.add_worksheet(title=sheet3_name, rows=100, cols=20)
            col1,col2 = st.columns([1.5,1.5])
            def update_cell_TitleMix_wikiFor(wikiFor):
                next_empty_row = helper.find_next_empty_cell_wikiFor_column(5)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 5, wikiFor)

            def update_cell_TitleMix_wikiBy(wikiBy):
                next_empty_row = helper.find_next_empty_cell_wikiBy_column(6)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 6, wikiBy)

            def update_cell_TitleMix_wikiArticleTitle(article):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleTitle_column(1)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 1, article)

            def update_cell_TitleMix_wikiArticleLink(url):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleLink_column(2)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 2, url)

            def update_cell_TitleMix_wikiArticleSummary(summary):
                next_empty_row = find_next_empty_cell_TitleMix_wikiArticleSummary_column(3)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 3, summary)

            def wikipedia_search(wikiArticleTitle):
                try:
                    page = wikipedia.page(wikiArticleTitle)
                    url = page.url
                    summary = page.summary
                    summary_sentences = sent_tokenize(summary)
                    summarized_summary = ' '.join(summary_sentences[:3])
                    return url, summarized_summary
                except wikipedia.exceptions.DisambiguationError as e:
                    options = e.options  # List of disambiguation options
                    first_option = options[0]  # Choose the first option automatically
                    try:
                        page = wikipedia.page(first_option)
                        url = page.url
                        summary = page.summary
                        summary_sentences = sent_tokenize(summary)
                        summarized_summary = ' '.join(summary_sentences[:3])
                        return url, summarized_summary
                    except wikipedia.exceptions.PageError:
                        st.warning(f"No Wikipedia page found for '{first_option}'")
                        return None, None
                except wikipedia.exceptions.PageError:
                    # If the page does not exist, return None for URL and summary
                    return None, None

            with col1:
                with st.form(key='info_form', clear_on_submit=True):
                    st.subheader("Info Gathering")
                    wikiNote = st.text_area("Note", height=208)
                    submitted = st.form_submit_button("Info Gathering")
                    sheet3_name = (wikiTitle + " WikiMix List")
                    if submitted:
                        helper.update_cell_wikiNote(wikiNote)
                        generate_new_wiki(wikiTitle)
                        wks_wikiTitle_data = database.worksheet(sheet3_name)
                        wks_wikiTitle_data.format('A1:F1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(1,4, 'WikiMix Note')
                        wks_wikiTitle_data.update_cell(2,4, wikiNote)

            with col2:
                #with st.form(key='search_form', clear_on_submit=True):
                
                st.subheader("Wiki Search")    
                
                import wikipediaapi

                sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
                sheet2_name = "data"
                cred_file = "creds.json"
                gc = gspread.service_account(cred_file)
                database = gc.open("WikiMix")
                wks_data = database.worksheet(sheet2_name)
                wiki_wiki = wikipediaapi.Wikipedia('MyProjectName (merlin@example.com)', 'en')
                #submitted4 = st.form_submit_button("Search")

                def wikipedia_search(keyword):
                    try:
                        page = wiki_wiki.page(keyword)  # Try to get the page based on the keyword
                        if page.exists():
                            url = page.fullurl
                            summary = page.summary
                            summary_sentences = sent_tokenize(summary)
                            summarized_summary = ' '.join(summary_sentences[:3])
                            return url, summarized_summary
                        else:
                            st.warning(f"No Wikipedia page found for '{keyword}'")
                            return None, None
                    except KeyError:
                        st.warning(f"No Wikipedia page found for '{keyword}'")
                        return None, None

                def find_next_empty_cell_wikiArticleTitle_column(column):
                    # Get all non-empty values in the specified column
                    values = wks_data.col_values(1)

                    # Find the index of the first empty cell
                    for idx, value in enumerate(values):
                        if not value:
                            return idx + 1  # Adding 1 to convert to 1-based index

                    # If no empty cell found, return the next index after the last cell
                    return len(values) + 1

                def update_cell_wikiArticleTitle(keyword):
                    next_empty_row = find_next_empty_cell_wikiArticleTitle_column(1)  # Column A is 1-based index
                    wks_data.update_cell(next_empty_row, 1, keyword)

                keyword = st.text_input("Enter Keyword")
                button = st.button("Search")
                if button:
                    url, summary = helper.wikipedia_search(keyword)
                    if url and summary:
                        st.markdown(f"**Search Result**")
                        st.markdown(f"{keyword}") 
                        st.markdown(f"**URL:** [Link]({url})")
                        st.markdown("**Summary:**")
                        st.write(summary)
                btn = st.button("Add Article")
                if btn:
                    update_cell_wikiArticleTitle(keyword)
                    page = wiki_wiki.page(keyword)
                    url = page.fullurl
                    summary = page.summary
                    helper.update_cell_wikiArticleLink(url)
                    helper.update_cell_wikiArticleSummary(summary)
                    st.success("Article Added. Go back and search for more articles!")          

            if wikiTitle and wikiFor and wikiBy and wikiNote:
                st.markdown("---")
                st.subheader("The Mix")
                sheet3_name = wikiTitle + " WikiMix List"
                wks_wikiTitle_data = database.worksheet(sheet3_name)
                article_list = wks_wikiTitle_data.col_values(1)[1:]

                def create_pdf_document(article_list):
                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    
                    pdf.set_font("Arial", size=28)
                    pdf.cell(200, 10, txt=wikiTitle + " WikiMix", ln=True, align="C")

                    pdf.set_font("Arial", size=14, style="B")
                    pdf.cell(200, 10, txt = "Note: ", ln=True, align='L')
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0,5, txt = wikiNote, align='L')
                    pdf.ln()

                    for article in article_list:
                        url, summary = helper.wikipedia_search(article)
                        if url and summary:
                            pdf.set_font("Arial", size=14, style="B")
                            pdf.cell(200, 10, txt=article.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                            pdf.set_text_color(0, 0, 255)
                            pdf.set_font("Arial", size=10)
                            pdf.cell(0, 10, txt=url, ln=True, link=url)
                            pdf.set_font("Arial", size=12, style="B")
                            pdf.set_text_color(0, 0, 0)
                            pdf.cell(200, 10, txt="Summary:", ln=True)
                            pdf.set_font("Arial", size=12)
                            pdf.multi_cell(0, 10, txt=summary.encode('latin-1', 'replace').decode('latin-1'),  align='L')
                            pdf.ln()
                    
                    return pdf

                def create_word_document(article_list):
                    doc = Document()
                    doc.add_heading(wikiTitle + " WikiMix", level=1)
                    doc.add_heading("Note:", level=2)
                    doc.add_paragraph(wikiNote)

                    bold_style = doc.styles.add_style('BoldStyle', WD_STYLE_TYPE.CHARACTER)
                    bold_font = bold_style.font
                    bold_font.bold = True
                    
                    for article in article_list:
                        url, summary = helper.wikipedia_search(article)
                        if url and summary:
                            doc.add_heading(article, level=2)
                            doc.add_heading(f"URL:", level=2).bold = True
                            doc.add_paragraph(f"{url}", style=None)
                            doc.add_heading(f"Summary:", level=2).bold = True
                            doc.add_paragraph(f"{summary}", style=None)
                    
                    return doc 

                if st.button("See "+ wikiTitle + " WikiMix"):
                    st.write("WikiMix: " + wikiTitle)
                    st.write("WikiMix Note: " + wks_wikiTitle_data.acell('D2').value)

                    for article in article_list:
                        if article:  # Check if the cell is not empty
                            url, summary = helper.wikipedia_search(article)
                            if url and summary:
                                st.markdown(f"**Title:** {article}")
                                st.markdown(f"**URL:** [Link]({url})")
                                st.markdown("**Summary:**")
                                st.write(summary)
                            else:
                                st.warning(f"No Wikipedia page found for '{article}'")
                        else:
                            break  # Break the loop when an empty cell is encountered

                    doc = create_word_document(article_list)
                    doc_filename = wikiTitle + " WikiMix.docx"
                    doc.save(doc_filename)
                    st.markdown(helper.get_download_link(doc_filename, "Download Doc"), unsafe_allow_html=True)
                    
                    pdf = create_pdf_document(article_list)
                    pdf_path = wikiTitle + " WikiMix.pdf"
                    create_pdf_document(article_list).output(pdf_path)
                    st.markdown(helper.get_download_link(pdf_path, "Download PDF"), unsafe_allow_html=True)

                st.markdown("---")
                # Email File
                st.subheader("Send WikiMix")
                email_sender = 'WikiMixList@gmail.com'
                email_password = 'ufpkegvfsqciasab'
                email_receiver = st.text_input("Email To")
                subject = "Here's your WikiMix"
                body = "Welcome to WikiMix, the app that let's you expertly curate a collection of Wikipedia articles that traverse diverse topics to offer an immersive journey through the realm of information." 
                
                pdf = create_pdf_document(article_list)
                pdf_path = wikiTitle + " WikiMix.pdf"
                create_pdf_document(article_list).output(pdf_path)
                
                em = EmailMessage()
                em['From'] = email_sender
                em['To'] = email_receiver
                em['Subject'] = subject
                em.set_content(body)

                context = ssl.create_default_context()
                button = st.button("Send WikiMix")
                filename = pdf_path

                if button:
                    with open(filename, 'rb') as f:
                        file_data = f.read()
                        em.add_attachment(file_data, maintype='application', subtype='pdf', filename=filename)
                    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
                        smtp.login(email_sender, email_password)
                        smtp.sendmail(email_sender,email_receiver, em.as_string())
                    st.success("WikiMix has been sent!")

    #RandomWikiMix
    elif menu_id == "Random WikiMix":
        st.header("Create a Random WikiMix")
        with st.container():
            sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
            sheet_name = "info_data"
            sheet2_name = "data"
            cred_file = "creds.json"
            gc = gspread.service_account(cred_file)
            database = gc.open("WikiMix")
            wks_info_data = database.worksheet(sheet_name)
            wks_data = database.worksheet(sheet2_name)
            wikiTitle = st.text_input("Title of WikiMix")
            if st.button('Submit WikiTitle'):
                helper.update_cell_wikiTitle(wikiTitle)
            def generate_new_wiki(wikiTitle):
                if wikiTitle != wks_info_data.col_values(1):
                    sheet3_name = (wikiTitle + " WikiMix List")
                    database.add_worksheet(title=sheet3_name, rows=100, cols=20)
            col1,col2 = st.columns([2,1])
            def update_cell_TitleMix_wikiFor(wikiFor):
                next_empty_row = helper.find_next_empty_cell_wikiFor_column(5)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 5, wikiFor)

            def update_cell_TitleMix_wikiBy(wikiBy):
                next_empty_row = helper.find_next_empty_cell_wikiBy_column(6)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 6, wikiBy)

            def update_cell_TitleMix_wikiArticleTitle(article):
                next_empty_row = helper.find_next_empty_cell_wikiArticleTitle_column(1)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 1, article)

            def update_cell_TitleMix_wikiArticleLink(url):
                next_empty_row = fhelper.ind_next_empty_cell_wikiArticleLink_column(2)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 2, url)

            def update_cell_TitleMix_wikiArticleSummary(summary):
                next_empty_row = helper.find_next_empty_cell_wikiArticleSummary_column(3)  # Column A is 1-based index
                wks_wikiTitle_data.update_cell(next_empty_row, 3, summary)

            with col1:
                with st.form(key='info_form', clear_on_submit=True):
                    st.subheader("Info Gathering")
                    wikiNote = st.text_area("Note", height=208)
                    submitted = st.form_submit_button("Info Gathering")
                    sheet3_name = (wikiTitle + " WikiMix List")
                    if submitted:
                        helper.update_cell_wikiNote(wikiNote)
                        generate_new_wiki(wikiTitle)
                        wks_wikiTitle_data = database.worksheet(sheet3_name)
                        wks_wikiTitle_data.update_cell(1,4, 'WikiMix Note')
                        wks_wikiTitle_data.format('D1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(2,4, wikiNote)
            
            with col2:
                with st.form(key='random_article', clear_on_submit=True):
                    st.subheader("Find Random Articles")
                    wikiTrackNumber = st.text_input("Number of articles for WikiMix")
                    wikiFor = st.text_input("Who is this WikiMix for?")
                    wikiBy = st.text_input("Who was this WikiMix made by?")
                    submitted2a = st.form_submit_button("Find Articles")
                    if submitted2a:
                        sheet3_name = (wikiTitle + " WikiMix List")
                        wks_wikiTitle_data = database.worksheet(sheet3_name)
                        wks_wikiTitle_data.format('A1:F1', {'textFormat': {'bold': True}})

                        wks_wikiTitle_data.update_cell(1,4, 'WikiMix Note')
                        wks_wikiTitle_data.format('D1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(2,4, wikiNote)

                        wks_wikiTitle_data.update_cell(1,5, 'WikiMix For')
                        wks_wikiTitle_data.format('E1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(2,5, wikiFor)

                        wks_wikiTitle_data.update_cell(1,6, 'WikiMix By')
                        wks_wikiTitle_data.format('F1', {'textFormat': {'bold': True}})
                        wks_wikiTitle_data.update_cell(2,6, wikiBy)
                        
                        update_cell_TitleMix_wikiFor(wikiFor)
                        update_cell_TitleMix_wikiBy(wikiBy)
                        helper.update_cell_wikiFor(wikiFor)
                        helper.update_cell_wikiBy(wikiBy)

            if wikiTitle and wikiFor and wikiBy and wikiNote:
                st.markdown("---")
                st.subheader("The Mix")
                if st.button("See "+ wikiTitle + " WikiMix"):
                    st.write("WikiMix: " + wikiTitle)
                    wks_wikiTitle_data = database.worksheet(sheet3_name)
                    st.write("WikiMix For: " + wks_wikiTitle_data.acell('E2').value)
                    st.write("WikiMix By: " + wks_wikiTitle_data.acell('F2').value)
                    st.write("WikiMix Note: " + wks_wikiTitle_data.acell('D2').value)
                    if wikiTrackNumber:
                        num_articles = int(wikiTrackNumber)
                        article_list = helper.randomWikiArticles(num_articles)
                        nltk.download('punkt')

                        def wikipedia_search(wikiSearch):
                            try:
                                page = wikipedia.page(wikiSearch)
                                url = page.url
                                summary = page.summary
                                summary_sentences = sent_tokenize(summary)  # Tokenize summary into sentences
                                summarized_summary = ' '.join(summary_sentences[:3])  # Join the first three sentences
                                return url, summarized_summary
                            except wikipedia.exceptions.DisambiguationError as e:
                                # If there are multiple matches, show the first one as an example
                                page = wikipedia.page(e.options[0])
                                url = page.url
                                summary = page.summary
                                summary_sentences = sent_tokenize(summary)
                                summarized_summary = ' '.join(summary_sentences[:3])
                                return url, summarized_summary
                            except wikipedia.exceptions.PageError:
                                # If the page does not exist, return None for URL and summary
                                return None, None

                        if article_list is not None and isinstance(article_list, list):
                            for article in article_list:
                                url, summary = helper.wikipedia_search(article)
                                if url and summary:
                                    st.markdown(f"**Title:** {article}")
                                    st.markdown(f"**URL:** [Link]({url})")
                                    st.markdown("**Summary:**")
                                    st.write(summary)
                                    st.markdown("---")
                                else:
                                    st.warning(f"No Wikipedia page found for '{article}'")
                        helper.update_cell_wikiArticleTitle(article)
                        helper.update_cell_wikiArticleLink(url)
                        helper.update_cell_wikiArticleSummary(summary) 
                        helper.create_word_document_random(article_list)
                        helper.create_pdf_document_random(article_list)
                        doc = helper.create_word_document_random(article_list)
                        doc_filename = wikiTitle + "WikiMix.docx"
                        doc.save(doc_filename)
                        st.markdown(helper.get_download_link(doc_filename, "Download Doc"), unsafe_allow_html=True)
                        pdf = helper.create_pdf_document_random(article_list)
                        pdf_path = wikiTitle + "WikiMix.pdf"
                        pdf.output(pdf_path)
                        st.markdown(helper.get_download_link(pdf_path, "Download PDF"), unsafe_allow_html=True)

    #ExploreWikiMix
    elif menu_id == "Explore WikiMix Database":
        st.header("Explore The WikiMix Database")
        col1,col2 = st.columns([2.25,.25])
        with st.container():
            with col1:
                # Connect to the Google Sheet
                sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
                sheet_name = "info_data"
                sheet2_name = "data"
                url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
                url2 = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet2_name}"
                cred_file = "creds.json"
                gc = gspread.service_account(cred_file)
                database = gc.open("WikiMix")
                wks_info_data = database.worksheet(sheet_name)
                wks_data = database.worksheet(sheet2_name)
                df = pd.read_csv(url, dtype=str, header=0, sep=",").fillna("")
                df2 = pd.read_csv(url2, dtype=str, header=0, sep=",").fillna("")

                #Think about adding Aggrid to this section for better sorting, filtering, and look options

                text_search = st.text_input("Search WikiMix Titles", value="").lower()
                # Filter the dataframe using masks
                if text_search:
                    m1 = df["WikiMix Title"].str.contains(text_search, case=False)
                    m2 = df["WikiMix For"].str.contains(text_search, case=False)
                    m3 = df["WikiMix By"].str.contains(text_search, case=False)
                    m4 = df["WikiMix Note"].str.contains(text_search, case=False)
                    df_search = df[m1 | m2 | m3 | m4]
                    df_reset = HTML(df_search.to_html(index=False))

                    # Show the results, if you have a text_search
                    if st.button("WikiMix Title Search"):
                        df_reset

                st.markdown("---")
                text_search2 = st.text_input("Search WikiMix Article Database", value="").lower()
                # Filter the dataframe using masks
                if text_search2:
                    m5 = df2["Wikipedia Article Title"].str.contains(text_search2, case=False)
                    m6 = df2["Wikipedia Article Link"].str.contains(text_search2, case=False)
                    m7 = df2["Wikipedia Article Summary"].str.contains(text_search2, case=False)
                    df_search2 = df[m5 | m6 | m7]
                    df_reset2 = HTML(df_search2.to_html(index=False))

                    # Show the results, if you have a text_search
                    if st.button("Article Search"):
                        df_reset2

    #Test Code Section
    elif menu_id == "Test":
        import wikipedia 
        import wikipediaapi
        
        wiki_wiki = wikipediaapi.Wikipedia('MyProjectName (merlin@example.com)', 'en')
        nlp = spacy.load("en_core_web_sm")

        def get_article_data(article_title):
            article = wiki_wiki.page(article_title)
            return article.text if article.exists() else None

        sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
        sheet4_name = "new_data"
        cred_file = "creds.json"
        gc = gspread.service_account(cred_file)
        database = gc.open("WikiMix")
        wks_new_data = database.worksheet(sheet4_name)

        # Function to get recommendations
        def get_recommendations(input_article, num_recommendations=10):
            input_text = get_article_data(input_article)
            if input_text:
                input_doc = nlp(input_text)
                article_titles = wikipedia.search(input_article, results=num_recommendations)
                
                recommendations = []
                for title in article_titles:
                    article_data = wiki_wiki.page(title)
                    if article_data.exists():
                        similarity = input_doc.similarity(nlp(article_data.text))
                        recommendations.append((title, article_data.fullurl, article_data.summary, similarity))
                
                recommendations.sort(key=lambda x: x[3], reverse=True)
                return recommendations

        def find_next_empty_cell_Rec_wikiArticleTitle_column(column):
            # Get all non-empty values in the specified column
            values = wks_new_data.col_values(1)

            # Find the index of the first empty cell
            for idx, value in enumerate(values):
                if not value:
                    return idx + 1  # Adding 1 to convert to 1-based index

            # If no empty cell found, return the next index after the last cell
            return len(values) + 1

        def find_next_empty_cell_Rec_wikiArticleLink_column(column):
            # Get all non-empty values in the specified column
            values = wks_new_data.col_values(2)

            # Find the index of the first empty cell
            for idx, value in enumerate(values):
                if not value:
                    return idx + 1  # Adding 1 to convert to 1-based index

            # If no empty cell found, return the next index after the last cell
            return len(values) + 1

        def find_next_empty_cell_Rec_wikiArticleSummary_column(column):
            # Get all non-empty values in the specified column
            values = wks_new_data.col_values(3)

            # Find the index of the first empty cell
            for idx, value in enumerate(values):
                if not value:
                    return idx + 1  # Adding 1 to convert to 1-based index

            # If no empty cell found, return the next index after the last cell
            return len(values) + 1

        def update_cell_Rec_wikiArticleTitle(article):
            next_empty_row = find_next_empty_cell_Rec_wikiArticleTitle_column(1)  # Column A is 1-based index
            wks_new_data.update_cell(next_empty_row, 1, article)

        def update_cell_Rec_wikiArticleLink(url):
            next_empty_row = find_next_empty_cell_Rec_wikiArticleLink_column(2)  # Column A is 1-based index
            wks_new_data.update_cell(next_empty_row, 2, url)

        def update_cell_Rec_wikiArticleSummary(summary):
            next_empty_row = find_next_empty_cell_Rec_wikiArticleSummary_column(3)  # Column A is 1-based index
            wks_new_data.update_cell(next_empty_row, 3, summary)

        # Streamlit UI
        st.title("Wikipedia Article Recommender")

        input_article = st.text_input("Enter an Article Title")
        if st.button("Find Recs"):
            if input_article:
                recommendations = get_recommendations(input_article, num_recommendations=2)
                
                if recommendations:
                    sheet_id = "1niITcc466j_tUjKEGRSU9wftuWM3qbboGM3XAj35dgg"
                    sheet4_name = "new_data"
                    cred_file = "creds.json"
                    gc = gspread.service_account(cred_file)
                    database = gc.open("WikiMix")
                    wks_new_data = database.worksheet(sheet4_name)
                    
                    st.subheader("Recommendations:")
                    for title, url, summary, similarity in recommendations:
                        st.markdown(f"**Title:** {title}")
                        st.markdown(f"**URL:** [Link]({url})")
                        st.markdown("**Summary:**")
                        st.write(summary)
                        st.markdown("---")
                        
                        #button_key = f"add_button_{url}"
                        #if st.button("Add", key=button_key):
                        update_cell_Rec_wikiArticleTitle(title)
                        update_cell_Rec_wikiArticleLink(url)
                        update_cell_Rec_wikiArticleSummary(summary)
                            #st.success(f"Added '{title}' to Google Sheets!")
                else:
                    st.warning("No recommendations found for the entered article.")


if __name__=="__main__":    
    main()
